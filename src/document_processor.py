"""
document_processor.py
Handles text extraction from PDF, DOCX, and EPUB files.
"""
import os
import re
import uuid
import chromadb
from chromadb.config import Settings
from chromadb.utils import embedding_functions

class DocumentProcessor:
    def __init__(self):
        self.pages = []          # list of {"index": int, "text": str, "label": str}
        self.current_page = 0
        self.file_path = None
        self.doc_type = None
        self.title = "Untitled Document"
        
        # ChromaDB setup
        self.chroma_client = chromadb.Client(Settings(anonymized_telemetry=False))
        # Use simple built-in default embeddings (all-MiniLM-L6-v2) for speed
        self.embedding_fn = embedding_functions.DefaultEmbeddingFunction()
        self.collection_name = "doc_collection"
        self.collection = None

    # ─────────────────────────── Public API ───────────────────────────

    def load(self, filepath: str) -> bool:
        """Load a document and return True on success."""
        ext = os.path.splitext(filepath)[1].lower()
        self.pages = []
        self.current_page = 0
        self.file_path = filepath
        self.title = os.path.basename(filepath)

        try:
            if ext == ".pdf":
                self._load_pdf(filepath)
                self.doc_type = "PDF"
            elif ext in (".docx", ".doc"):
                self._load_docx(filepath)
                self.doc_type = "DOCX"
            elif ext == ".epub":
                self._load_epub(filepath)
                self.doc_type = "EPUB"
            elif ext == ".txt":
                self._load_txt(filepath)
                self.doc_type = "TXT"
            else:
                return False
                
            if len(self.pages) > 0:
                self._build_vector_index()
                return True
            return False
        except Exception as e:
            print(f"[DocumentProcessor] Load error: {e}")
            return False

    def unload(self, delete_file: bool = False) -> bool:
        """Clear document from memory and optionally delete the file from disk."""
        try:
            if delete_file and self.file_path and os.path.isfile(self.file_path):
                try:
                    os.remove(self.file_path)
                except Exception as e:
                    print(f"Could not delete file {self.file_path}: {e}")

            if self.collection:
                try:
                    self.chroma_client.delete_collection(name=self.collection_name)
                    self.collection = None
                except Exception:
                    pass

            self.pages = []
            self.current_page = 0
            self.file_path = None
            self.doc_type = None
            self.title = "Untitled Document"
            return True
        except Exception as e:
            print(f"[DocumentProcessor] Unload error: {e}")
            return False

    def page_count(self) -> int:
        return len(self.pages)

    def get_page(self, index: int) -> str:
        if 0 <= index < len(self.pages):
            return self.pages[index]["text"]
        return ""

    def get_current_text(self) -> str:
        return self.get_page(self.current_page)

    def get_current_label(self) -> str:
        if 0 <= self.current_page < len(self.pages):
            return self.pages[self.current_page]["label"]
        return "Unknown"

    def next_page(self) -> bool:
        if self.current_page < len(self.pages) - 1:
            self.current_page += 1
            return True
        return False

    def prev_page(self) -> bool:
        if self.current_page > 0:
            self.current_page -= 1
            return True
        return False

    def go_to_page(self, index: int) -> bool:
        if 0 <= index < len(self.pages):
            self.current_page = index
            return True
        return False

    def get_full_text(self, max_chars: int = 50000) -> str:
        """Return combined text of all pages (truncated for AI)."""
        combined = "\n\n".join(p["text"] for p in self.pages)
        return combined[:max_chars]

    def get_chapter_text(self, chapter_num: int) -> str:
        """Return text of a specific chapter/page (1-indexed)."""
        idx = chapter_num - 1
        if 0 <= idx < len(self.pages):
            return self.pages[idx]["text"]
        return ""

    def search(self, query: str) -> list:
        """Search text and return list of (page_index, snippet) matches."""
        results = []
        q = query.lower()
        for page in self.pages:
            text = page["text"].lower()
            pos = text.find(q)
            if pos != -1:
                snippet = page["text"][max(0, pos-60):pos+120]
                results.append({"page": page["index"], "label": page["label"], "snippet": snippet})
        return results

    def get_relevant_context(self, query: str, n_results: int = 4) -> str:
        """Fetch the most relevant text chunks for a given query via ChromaDB."""
        if not self.collection or self.collection.count() == 0:
            return self.get_full_text(max_chars=10000)

        # Ensure we don't ask for more results than we have chunks
        k = min(n_results, self.collection.count())
        if k == 0:
            return ""
            
        try:
            results = self.collection.query(
                query_texts=[query],
                n_results=k
            )
            
            if not results["documents"] or not results["documents"][0]:
                return self.get_full_text(max_chars=10000)
                
            # Combine the returned chunks
            documents = results["documents"][0]
            context = "\n...\n".join(documents)
            return context
        except Exception as e:
            print(f"[DocumentProcessor] Vector search failed: {e}. Falling back to full text.")
            return self.get_full_text(max_chars=10000)

    # ─────────────────────────── Private loaders ───────────────────────

    def _build_vector_index(self):
        """Index all loaded pages into an ephemeral Chroma collection."""
        try:
            # Recreate collection to clear old data
            try:
                self.chroma_client.delete_collection(name=self.collection_name)
            except Exception:
                pass
                
            self.collection = self.chroma_client.create_collection(
                name=self.collection_name, 
                embedding_function=self.embedding_fn
            )
            
            # Simple chunking logic to ensure we don't hit max payload sizes
            docs = []
            metadatas = []
            ids = []
            
            for page in self.pages:
                text = page["text"]
                # Chunk aggressively for better vector retrieval
                words = text.split()
                chunk_size = 300
                overlap = 50
                
                if not words: continue
                
                for i in range(0, len(words), chunk_size - overlap):
                    chunk = " ".join(words[i:i + chunk_size])
                    docs.append(chunk)
                    metadatas.append({"page": page["index"], "label": page["label"]})
                    ids.append(str(uuid.uuid4()))
                    
            if docs:
                self.collection.add(
                    documents=docs,
                    metadatas=metadatas,
                    ids=ids
                )
                print(f"[DocumentProcessor] Indexed {len(docs)} chunks into ChromaDB.")
        except Exception as e:
            print(f"[DocumentProcessor] Error bulding vector index: {e}")
            self.collection = None

    def _load_pdf(self, filepath: str):
        import fitz  # PyMuPDF
        doc = fitz.open(filepath)
        for i, page in enumerate(doc):
            text = page.get_text("text").strip()
            if text:
                self.pages.append({
                    "index": i,
                    "text": text,
                    "label": f"Page {i + 1}"
                })
        doc.close()

    def _load_docx(self, filepath: str):
        from docx import Document
        doc = Document(filepath)
        # Split by headings into chapters, otherwise into chunks
        current_chunk = []
        chapter_idx = 0
        chapter_label = "Section 1"

        for para in doc.paragraphs:
            if para.style.name.startswith("Heading"):
                if current_chunk:
                    self.pages.append({
                        "index": chapter_idx,
                        "text": "\n".join(current_chunk),
                        "label": chapter_label
                    })
                    chapter_idx += 1
                chapter_label = para.text.strip() or f"Section {chapter_idx + 1}"
                current_chunk = []
            else:
                if para.text.strip():
                    current_chunk.append(para.text.strip())

        if current_chunk:
            self.pages.append({
                "index": chapter_idx,
                "text": "\n".join(current_chunk),
                "label": chapter_label
            })

        # If no headings found, chunk by ~500 words
        if not self.pages:
            all_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            words = all_text.split()
            chunk_size = 500
            for i in range(0, len(words), chunk_size):
                chunk = " ".join(words[i:i + chunk_size])
                self.pages.append({
                    "index": i // chunk_size,
                    "text": chunk,
                    "label": f"Section {i // chunk_size + 1}"
                })

    def _load_epub(self, filepath: str):
        import ebooklib
        from ebooklib import epub
        from bs4 import BeautifulSoup

        book = epub.read_epub(filepath)
        chapter_idx = 0
        for item in book.get_items():
            if item.get_type() == ebooklib.ITEM_DOCUMENT:
                soup = BeautifulSoup(item.get_content(), "html.parser")
                text = soup.get_text(separator="\n").strip()
                if len(text) > 100:
                    title_tag = soup.find(["h1", "h2", "h3"])
                    label = title_tag.get_text().strip() if title_tag else f"Chapter {chapter_idx + 1}"
                    self.pages.append({
                        "index": chapter_idx,
                        "text": text,
                        "label": label
                    })
                    chapter_idx += 1

    def _load_txt(self, filepath: str):
        with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()
        words = content.split()
        chunk_size = 600
        for i in range(0, len(words), chunk_size):
            chunk = " ".join(words[i:i + chunk_size])
            self.pages.append({
                "index": i // chunk_size,
                "text": chunk,
                "label": f"Section {i // chunk_size + 1}"
            })
